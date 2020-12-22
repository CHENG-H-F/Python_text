import base64
import re
import sys
import uuid
from operator import itemgetter

from django.http import HttpResponse, JsonResponse, FileResponse, Http404
from django.shortcuts import render
from django.views.decorators import csrf
import numpy as np
import os
import json
import smtplib
import inspect
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .DataRequest import getData
from docx.text.paragraph import Paragraph
from docx.styles import style
from docx import Document
from docx.document import Document as Doc
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml.ns import qn
from .sql import sql
from django.views.decorators.csrf import csrf_exempt


def index(request):
    return render(request, 'index.html')


def top(request):
    return render(request, 'top.html')


def menu(request):
    return render(request, 'menu.html')


def default(request):
    return render(request, 'default.html')


def status_page(request):
    col = sql.returnAllData(request)
    return render(request, 'status_page.html', {'col': col})


def fetch(request):
    sql.insertAll(request)
    return status_page(request)


def singleLine_page(request):
    col = sql.returnAllData(request)
    return render(request, 'singleLine_page.html', {'col': col})


def search(request):
    global name
    if request.method == "POST":
        name = request.POST['name']
    data = sql.return_num_and_time(name, request)
    return HttpResponse(json.dumps(data))


def compare_page(request):
    col = sql.returnAllData(request)
    return render(request, 'compare_page.html', {'col': col})


#  根据多个name获取记录
def compare(request):
    global na, data
    if request.method == "POST":
        na = request.POST  # <QueryDict: {'names[]': ['上海欢乐谷', '大宁郁金香公园', '上海召稼楼景区']}>
    na = dict(na)  # {'names[]': ['上海欢乐谷', '大宁郁金香公园', '上海召稼楼景区']}
    names = na.get('names[]', None)  # ['上海欢乐谷', '大宁郁金香公园', '上海召稼楼景区']
    series = []
    content = ''
    if (names != None):
        for i in names:
            data = sql.return_num_and_time(i, request)  # 根据名字获取记录，并循环。
            newitem = {'data': data['num'], 'name': i, 'type': 'line', 'smooth': 'true',
                       'areaStyle': {'opacity': '0.7', }}  # 每一个name对应的都取最新的一条创建新json
            series.append(newitem)
        content = {"date": data['time'], "name": names, "series": series}
    return HttpResponse(json.dumps(content))


def report_page(request):
    col = sql.returnAllData(request)
    return render(request, 'report_page.html', {'col': col})


def report(request):
    global na_pic_txt, data
    if request.method == "POST":
        na_pic_txt = request.POST  # <QueryDict: {'names[]': ['上海欢乐谷', '大宁郁金香公园', '上海召稼楼景区']}>
    na_pic_txt = dict(na_pic_txt)  # {'names[]': ['上海欢乐谷', '大宁郁金香公园', '上海召稼楼景区']}
    names = na_pic_txt.get('names[]', None)  # ['上海欢乐谷', '大宁郁金香公园', '上海召稼楼景区']
    series = []
    content = ''
    if (names != None):
        for i in names:
            data = sql.return_num_and_time(i, request)  # 根据名字获取记录，并循环。
            newitem = {'data': data['num'], 'name': i, 'type': 'line', 'smooth': 'true',
                       'areaStyle': {'opacity': '0.7', }}  # 每一个name对应的都取最新的一条创建新json
            series.append(newitem)
        content = {"date": data['time'], "name": names, "series": series}
    return HttpResponse(json.dumps(content))


def create_img(request):
    global na_pic_txt
    if request.method == "POST":
        na_pic_txt = request.POST
    na_pic_txt = dict(na_pic_txt)
    picInfo = na_pic_txt.get('picInfo', None)[0]
    name = na_pic_txt.get('name', None)[0]
    if (picInfo):
        picInfo = picInfo.replace('data:image/png;base64,', '')  # 将base64前部分的图片标识去除掉
        imgdata = base64.b64decode(picInfo)  # 对base64解码，生成img的二进制文件
        # 生成一个随机字符串
        uuid_str = uuid.uuid4().hex
        file_name = uuid_str + '_' + name
        # 构成完整文件存储路径
        tmp_file_name = 'static/css/images/echarts_img/' + file_name + '.jpg'
        file = open(tmp_file_name, 'wb')
        request.session[file_name] = tmp_file_name  # 将文件路径储存在session中，便于下载报告。
        file.write(imgdata)
        file.close()
        if (name == 'Chart with mutiple spot data'):
            return HttpResponse(2)
        else:
            return HttpResponse(1)
    else:
        return HttpResponse(0)


def download_report(request):
    global file_title
    if request.method == "POST":
        file_title = request.POST['file_title']
    list = dict(request.session)
    if len(list) != 0:
        filename = createReport(request, list, file_title)
        return HttpResponse(1)
    else:
        return HttpResponse(0)


def clear_session(request):
    request.session.clear()
    return HttpResponse(1)


def createReport(request, list, file_title):
    # create a document
    document: Doc = Document()
    # set built-in template style
    doc_style: style._ParagraphStyle = document.styles['Normal']
    doc_style.font.name = u'宋体'
    doc_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # search data from databases
    name_list = []
    name = []
    databases_list = []
    for i in list.keys():
        name_list.append(re.findall('_(.*)', i)[0])
    for i in range(0, len(name_list) - 1):
        d_list = sql.returnDataByName(request, name_list[i])
        if (len(d_list) >= 1):
            d_list = d_list[-1:]
            databases_list.append(d_list.pop(0))
    databases_list = sorted(databases_list, key=itemgetter('NUM'), reverse=True)
    # create headings
    document.add_heading(file_title, level=1)
    document.add_heading('Attractions', level=2)
    # create table
    document.add_paragraph('Attraction Name List').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    len_list = len(list)
    table = document.add_table(rows=len_list, cols=7, style="Light List Accent 2")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Code'
    hdr_cells[2].text = 'Name'
    hdr_cells[3].text = 'Address'
    hdr_cells[4].text = 'Grade'
    hdr_cells[5].text = 'Business Hours'
    hdr_cells[6].text = 'Recent Flow'
    for i in range(1, len_list):
        hdr_cells = table.rows[i].cells
        hdr_cells[0].text = str(i)
        hdr_cells[1].text = databases_list[i - 1]['CODE']
        db_NAME = databases_list[i - 1]['NAME']
        name.append(db_NAME)
        hdr_cells[2].text = db_NAME
        hdr_cells[3].text = databases_list[i - 1]['ADDRESS']
        hdr_cells[4].text = databases_list[i - 1]['GRADE']
        hdr_cells[5].text = databases_list[i - 1]['T_TIME']
        hdr_cells[6].text = str(databases_list[i - 1]['NUM'])
    name_list.sort(reverse=True)
    document.add_heading('Chart with single data', level=2)
    name.append('Chart with mutiple spot data')
    for i in name:
        document.add_paragraph('Visitor flow chart of ' + i).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for j in list.keys():
            if (i == re.findall('_(.*)', j)[0]):
                if (i == 'Chart with mutiple spot data'):
                    document.add_heading('Multi-attraction data comparison', level=2)
                document.add_picture(list[j], width=Inches(6.5))
    filename = 'static/css/Documents/' + file_title + '.docx'
    document.save(filename)
    return filename


def download_template(request):
    global file_title
    if request.method == "POST":
        file_title = request.POST.get('file_title', None)
    file_directory = 'static/css/Documents/' + file_title + '.docx'
    file = open(file_directory, 'rb')
    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream;charset=UTF-8;'
    # response['Content-Disposition'] = 'attachment;filename="' + file_title + '.docx"'
    str = 'attachment;filename="'+ file_title + '.docx"'
    response['Content-Disposition'] = str.encode('utf-8')
    return response


def email_page(request):
    filename = []
    file_directory = 'static/css/Documents'
    filename = os.listdir(file_directory)
    return render(request, 'email_page.html', {'filename': filename})


def send_email(request):
    global re
    if request.method == "POST":
        re = request.POST
    re_list = dict(re)
    email_address = re_list.get('email_address', None)[0]  # 获取邮件地址
    email_address = email_address.split(';')
    file_name = re_list.get('file_name[]', None)  # 获取选择的文件名

    # 发送邮件
    server_address = 'smtp.163.com'
    to_who = email_address
    from_who = '15007939247@163.com'
    from_who_password = 'HHUDODBBLBWQUOGR'

    subjext = '这是一封很有意思的邮件哦！'

    content = '''
    来自:%(from_who)s<br>
    发往:%(to_who)s<br>
    主题:这是一封很有意思的邮件哦！<br>
    <br>
    你好，这是我的一封邮件,它带有附件哟!
    ''' % {'to_who': ','.join(to_who), 'from_who': from_who}
    content_part = MIMEText(content, 'html', 'utf-8')

    msg = MIMEMultipart()
    msg["Accept-Charset"] = "ISO-8859-1,utf-8"
    msg["Accept-Language"] = "zh-CN"
    msg['From'] = '15007939247@163.com'
    msg['To'] = ','.join(email_address)
    msg['Subject'] = subjext
    msg.attach(content_part)

    # 如果有附件则上传附件
    if file_name:
        for i in file_name:
            attachment_part = MIMEApplication(open('static/css/Documents/' + i, 'rb').read())
            attachment_part["Content-Type"] = 'application/octet-stream'
            attachment_part.add_header('Content-Disposition', 'attachment', filename=i)
            msg.attach(attachment_part)
    try:
        server = smtplib.SMTP_SSL(server_address, 465)
        server.login(from_who, from_who_password)
        server.sendmail(from_who, to_who, msg.as_string())
        server.quit()
        return HttpResponse(1)
    except smtplib.SMTPException as e:
        print(e)
        return HttpResponse(0)
